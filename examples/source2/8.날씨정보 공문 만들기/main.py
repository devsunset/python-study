from docx import Document
import docx
from docx.oxml.ns import qn
from datetime import date
import requests
import re

import os
os.chdir(os.path.dirname(os.path.abspath(__file__)))

#날씨정보 읽어오기
url = "http://www.kma.go.kr/wid/queryDFSRSS.jsp?zone=4139054000"
response = requests.get(url)

요청시간 = re.findall(r'<pubDate>(.+)</pubDate>',response.text)
지역 = re.findall(r'<category>(.+)</category>',response.text)
날씨정보 = re.findall(r'<wfKor>(.+)</wfKor>',response.text)
시간 = re.findall(r'<hour>(.+)</hour>',response.text)
온도 = re.findall(r'<temp>(.+)</temp>',response.text)
습도 = re.findall(r'<reh>(.+)</reh>',response.text)

print("요청시간: ",요청시간)
print("지역: ",지역)
print("날씨정보: ",날씨정보)
print("시간: ",시간)
print("온도: ",온도)
print("습도: ",습도)

#공문양식에서 읽어 문서번호, 발신일자, 수신, 제목 작성
doc = Document('공문양식.docx')

run = doc.paragraphs[2].add_run('2022-0005')  #문서번호
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')

today = date.today()
run = doc.paragraphs[3].add_run(str(today)) #발신일자
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')

run = doc.paragraphs[4].add_run('땡땡회사') #수신
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')

run = doc.paragraphs[6].add_run('오늘의 날씨정보') #제목
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')

#표에 값 넣기
doc.tables[0].rows[0].cells[1].paragraphs[0].text = "지역"

for i in range(4):
    doc.tables[0].rows[1].cells[i+1].paragraphs[0].clear()
    run = doc.tables[0].rows[1].cells[i+1].paragraphs[0].add_run(str(today) + '\n' + 시간[i] + '시')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(12)
    
    doc.tables[0].rows[2].cells[i+1].paragraphs[0].clear()
    run = doc.tables[0].rows[2].cells[i+1].paragraphs[0].add_run(온도[i])
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(12)
    
    doc.tables[0].rows[3].cells[i+1].paragraphs[0].clear()
    run = doc.tables[0].rows[3].cells[i+1].paragraphs[0].add_run(날씨정보[i])
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(12)

doc.save('공문생성.docx')