from docx import Document
import docx
from docx.oxml.ns import qn
import os
os.chdir(os.path.dirname(os.path.abspath(__file__)))

doc = Document('공문양식.docx')

for i, paragraph in enumerate(doc.paragraphs):
    print(str(i) + ": " + paragraph.text)

table = doc.tables[0]
for row in table.rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            print(para.text)