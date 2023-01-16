from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

customers = [
    'A고객사', 'B고객사', 'C고객사', 'D고객사',
    'E고객사', 'F고객사', 'G고객사', 'H고객사',
]

for (index, customer) in enumerate(customers):
    document = Document("template.docx")
    paras = document.paragraphs

    paras[0].text = paras[0].text.replace('{company}', '주식회사 파이썬')

    content = paras[2].text
    content = content.replace('{number}', f'ABCD-1234-{str(index).zfill(4)}')
    content = content.replace('{date}', '2022-12-12')
    content = content.replace('{receiver}', customer)
    content = content.replace('{sender}', '주식회사 파이썬')
    content = content.replace('{title}', '사내 워크숍 진행으로 인한 업무 안내')
    paras[2].text = content

    document.add_paragraph('''주식회사 파이썬 사내 워크숍을 진행합니다.
이로 인해 12월 12일(월), 12월 13일(화) 2일간 전화 연결이 어려울 수 있습니다.
문의사항이 있다면 python@python-freelec-book.com으로 보내주시면
회신 드리도록 하겠습니다.
감사합니다.''')
    p = document.add_paragraph(style='List Number')
    p.add_run('일시\t: ').bold = True
    p.add_run('2022-12-12 ~ 2022-12-13')

    p = document.add_paragraph(style='List Number')
    p.add_run('긴급연락\t: ').bold = True
    p.add_run('010-1234-5678')

    document.add_paragraph('')
    p = document.add_paragraph('주식회사 파이썬 임직원 일동')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.save(f"{customer}.docx")